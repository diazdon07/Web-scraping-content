import xml.etree.ElementTree as ET
from docx import Document
from bs4 import BeautifulSoup
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext

def process_export():
    try:
        xml_path = xml_file_path.get()
        output_path = output_file_path.get()
        urls_text = urls_box.get("1.0", tk.END).strip()
        urls = [u.strip().rstrip("/") for u in urls_text.splitlines() if u.strip()]

        if not xml_path:
            messagebox.showerror("Error", "Please attach your WordPress XML export file.")
            return
        if not urls:
            messagebox.showerror("Error", "Please enter at least one URL.")
            return
        if not output_path:
            output_path = "wp-export-structured.docx"

        # Namespaces
        ns = {
            "content": "http://purl.org/rss/1.0/modules/content/",
            "dc": "http://purl.org/dc/elements/1.1/",
            "wp": "http://wordpress.org/export/1.2/",
        }

        # Parse XML
        tree = ET.parse(xml_path)
        root = tree.getroot()

        doc = Document()
        doc.add_heading("WordPress Export – Selected Site Content", level=0)

        for item in root.findall("./channel/item"):
            link = item.find("link")
            if link is not None:
                url = link.text.strip().rstrip("/")
                if url in urls:
                    slug = url.split("://")[-1].split("/", 1)[-1]  # everything after domain
                    title = item.find("title").text if item.find("title") is not None else "Untitled"
                    meta_title = title

                    # Extract content
                    content_encoded = item.find("content:encoded", ns)
                    content_html = content_encoded.text if content_encoded is not None else ""
                    soup = BeautifulSoup(content_html, "html.parser")

                    # Page Header
                    doc.add_page_break()
                    doc.add_heading(f"Page: {title}", level=1)
                    doc.add_paragraph(f"Slug: /{slug}")
                    doc.add_paragraph(f"Page Meta Title: {meta_title}")
                    doc.add_paragraph("Page Meta Description:")  # blank
                    doc.add_paragraph("")

                    # Structured content
                    for elem in soup.descendants:
                        if elem.name:
                            if elem.name == "h1":
                                doc.add_heading(elem.get_text(strip=True), level=1)
                            elif elem.name == "h2":
                                doc.add_heading(elem.get_text(strip=True), level=2)
                            elif elem.name == "h3":
                                doc.add_heading(elem.get_text(strip=True), level=3)
                            elif elem.name == "p":
                                text = elem.get_text(strip=True)
                                if text:
                                    doc.add_paragraph(text)
                            elif elem.name in ["ul", "ol"]:
                                for li in elem.find_all("li"):
                                    doc.add_paragraph(li.get_text(strip=True), style="List Bullet")
                            elif elem.name == "blockquote":
                                text = elem.get_text(strip=True)
                                if text:
                                    doc.add_paragraph(text, style="Intense Quote")

        # Save file
        doc.save(output_path)
        messagebox.showinfo("Success", f"✅ File saved as {output_path}")

    except Exception as e:
        messagebox.showerror("Error", str(e))

def browse_xml():
    file_path = filedialog.askopenfilename(filetypes=[("WordPress Export", "*.xml")])
    if file_path:
        xml_file_path.set(file_path)
        attached_label.config(text=f"📂 Attached: {file_path}")

def save_as_docx():
    file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                             filetypes=[("Word Document", "*.docx")])
    if file_path:
        output_file_path.set(file_path)

# --- GUI Window ---
root = tk.Tk()
root.title("WordPress Export → DOCX Converter")
root.geometry("700x650")

xml_file_path = tk.StringVar()
output_file_path = tk.StringVar()

# XML File Picker
tk.Label(root, text="Attach WordPress XML Export File:").pack(anchor="w", padx=10, pady=5)
tk.Button(root, text="Attach File", command=browse_xml).pack(padx=10, pady=5)
attached_label = tk.Label(root, text="📂 No file attached", fg="blue")
attached_label.pack(anchor="w", padx=10)

# URLs Box
tk.Label(root, text="Enter URLs (one per line):").pack(anchor="w", padx=10, pady=5)
urls_box = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=80, height=15)
urls_box.pack(padx=10, pady=2)

# Process Button
tk.Button(root, text="Generate DOCX", command=process_export,
          bg="green", fg="white", font=("Arial", 12, "bold")).pack(pady=20)

root.mainloop()